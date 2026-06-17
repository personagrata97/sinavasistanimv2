#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Sendikasyon alıcı kredisi – tam çalışma kağıdı (nihai mail ile uyumlu)."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "Sendikasyon_Alici_Kredisi_Calisma_Kagidi.docx"


def h(doc, text, level=2):
    doc.add_heading(text, level=level)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def b(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    for run in para.runs:
        run.font.size = Pt(11)


def t(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, hd in enumerate(headers):
        table.rows[0].cells[i].text = hd
        for r in table.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
            for r in table.rows[ri + 1].cells[ci].paragraphs[0].runs:
                r.font.size = Pt(10)
    doc.add_paragraph()


def build():
    doc = Document()
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    t0 = doc.add_paragraph()
    t0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t0.add_run("ÇALIŞMA KAĞIDI (TAM METİN)\nSendikasyon Alıcı Kredisi – Ajan Banka Ödeme Yapısı")
    r.bold = True
    r.font.size = Pt(14)
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = m.add_run("Türk Eximbank uyum görüşü desteği | Haziran 2026 | Nihai mail ile uyumlu")
    mr.font.size = Pt(10)
    doc.add_paragraph()

    # A
    h(doc, "A. BU BELGE NE İÇİN?")
    p(doc, "İç birimden gelen soruya verilen mevzuat görüşünün arkasındaki mantığı, dayanakları ve baskıda söylenecek cevapları tek yerde toplar.")
    p(doc, "Kapsam: Sermaye Hareketleri Genelgesi + 32 sayılı Karar. Eximbank iç prosedür, BDDK, OECD ayrı dosya.")

    # B
    h(doc, "B. KELİME SÖZLÜĞÜ (kısaltma yok)")
    t(doc, ["Kelime", "Ne demek?"], [
        ("Alıcı kredisi", "Yabancı alıcı (Tanzanya) Türkiye'den mal alıyor; ödemeyi vadeli/krediyle yapıyor; ihracatçı peşin para alıyor."),
        ("Sendikasyon", "Birden fazla bankanın aynı krediye ortak olması."),
        ("Ajan banka", "Sendikasyonda parayı toplayan, dağıtan, süreci yöneten banka. BU İŞTE: Standard Chartered."),
        ("Katılımcı banka", "Sendikasyona payıyla giren banka. BU İŞTE: Türk Eximbank (Bankamız)."),
        ("Kullandırım", "Krediden fiilen para çekilmesi; ödeme günü."),
        ("Katılım payı", "Sendikasyonda Eximbank'ın üstlendiği tutar."),
        ("İhracatçı / lehtar", "Malı satan, parayı alan. Yapı Merkezi."),
        ("Borçlu", "Krediyi geri ödeyecek taraf. Tanzanya devleti (anlaşmaya göre)."),
        ("Sermaye Hareketleri Genelgesi", "Döviz kredisi kullanımında bankaların uyacağı TCMB genelgesi."),
        ("32 sayılı Karar", "Döviz kredisi kullanımının ana yasal çerçevesi."),
        ("Risk Merkezi", "Bankalar Birliği kredi takip sistemi. Bildirim zorunlu (17. md.)."),
        ("İstatistik GM", "TCMB dış borç istatistiği. Sendikasyon katılım payında bildirim yok (30. md.)."),
        ("İhracat taahhüdü (2017/4)", "Firma döviz kredi kullanırken 'bu kadar ihracat yapacağım' taahhüdü. İhracat kredisinde. Alıcı kredisinde ihracatçıya DOĞMAZ."),
        ("SCB", "Standard Chartered Bank = ajan banka (başka kurum değil)."),
    ])

    # C
    h(doc, "C. İŞİN ÖZÜ – ADIM ADIM")
    b(doc, "Yapı Merkezi Tanzanya'da demiryolu projesi işi yapacak (ihracat/taahhüt işi).")
    b(doc, "Tanzanya hemen ödemeyecek; alıcı kredisi ile finanse edilecek.")
    b(doc, "Finansman sendikasyon: Eximbank + Standard Chartered + Afreximbank, DBSA, KUKE, EKN, SACE.")
    b(doc, "Kredi anlaşması: Eximbank, Standard Chartered ve Tanzanya devleti arasında.")
    b(doc, "Her kullandırımda: Eximbank katılım payını Standard Chartered'a gönderir.")
    b(doc, "Standard Chartered (ajan) tutarı Yapı Merkezi'nin Türkiye'deki banka hesabına öder.")
    b(doc, "Finansman işlemlerinde ödeme ihracatçı TR hesabına; refinansmanda farklı akış (bu mail finansman odaklı).")

    # D
    h(doc, "D. TARAFLAR – KARIŞTIRMA YAPMA")
    t(doc, ["Rol", "Kim", "Ne yapar?"], [
        ("Katılımcı – Bankamız", "Türk Eximbank", "Sendikasyona katılır; kullandırımda payını ajan bankaya aktarır. AJAN DEĞİL."),
        ("Ajan banka", "Standard Chartered Bank", "Parayı toplar/dağıtır; ihracatçıya öder. Yurt dışında yerleşik."),
        ("İhracatçı", "Yapı Merkezi", "Satış bedelini TR hesabında alır. Kredi kullanan değil."),
        ("Borçlu", "Tanzanya devleti", "Kredi borcunu (anlaşmaya göre) öder."),
        ("Destek kuruluşları", "KUKE, EKN, SACE vb.", "İhracat kredi/garanti kuruluşu; 27. md. bağlamı."),
    ])
    p(doc, "Ödeme akışı: Türk Eximbank → Standard Chartered (ajan) → Yapı Merkezi (TR hesabı)", bold=True)

    # E
    h(doc, "E. GELEN SORU NE?")
    p(doc, "İşlem birimi soruyor: 'Kullandırımda önce bizden Standard Chartered'a, sonra ajanın ihracatçı TR hesabına ödeme yapılması mevzuata aykırı mı, engel var mı?'")
    p(doc, "Malum: Alıcı kredisi finansmanında genelde ihracatçının Türkiye hesabına ödeme yapılıyor.")

    # F
    h(doc, "F. VERİLEN CEVAP (NİHAİ MAIL ÖZETİ)")
    b(doc, "Sonuç: Sermaye Hareketleri Genelgesi ve 32 sayılı Karar kapsamında engel teşkil etmez.")
    b(doc, "Dayanak maddeler: 30 (sendikasyon), 27 (alıcı kredisi), 19 (banka aracılığı), 23 (para yolu).")
    b(doc, "30(2) uygulanmaz: ajan yurt dışında (Standard Chartered).")
    b(doc, "Alıcı kredisinde 2017/4 ihracat taahhüdü doğmaz; taahhüt kapatma yok.")
    b(doc, "Yükümlülükler: 19, 14, 17, 30, 33 (ECA kaynaklı), 31 — taahhüt takibi YOK.")

    # G
    h(doc, "G. İLK VERSİYONDAN NE DEĞİŞTİ? (hata düzeltmeleri)")
    t(doc, ["Konu", "İlk (yanlış/fazla)", "Nihai (doğru)"], [
        ("İhracat taahhüdü", "Takibi yazıldı", "Doğmaz; mailde açık yazıldı; takip yükümlülüğü yok"),
        ("Ajan banka", "Bazen karışık", "Standard Chartered; Eximbank katılımcı"),
        ("Sonuç tonu", "Bazen belirsiz", "Engel teşkil etmez (SHG+32 Karar)"),
        ("Eximbank rolü", "Net değildi", "Katılımcı; payı ajan'a gönderen"),
    ])

    # H
    h(doc, "H. MEVZUAT MADDELERİ – TEK TEK")
    h(doc, "Madde 30 (1) – Sendikasyon", 3)
    p(doc, "Türk bankaları + yurt dışı bankalar sendikasyon; işlemler dövizle.")
    p(doc, "Türk bankasının katılım payı: yurt içi döviz kredisi gibi izlenir.")
    p(doc, "Katılım payı: dış finansman numarası YOK; İstatistik dış borç bildirimi YOK.")
    p(doc, "İşlemler yurt dışı girişimci banka eliyle; kullanım/geri ödeme yurt dışı kredi esasları.")
    p(doc, "Bu iş: Eximbank payını ajan (SCB) eliyle dağıtım → 30(1) çerçevesi.")

    h(doc, "Madde 30 (2) – Türkiye'de girişimci ajan", 3)
    p(doc, "Türkiye'de girişimci banka ajan; yabancı banka kaynaklı sendikasyon; ajan katılmazsa farklı rejim.")
    p(doc, "BU İŞTE UYGULANMAZ: ajan Standard Chartered (yurt dışı), Eximbank katılımcı.")

    h(doc, "Madde 27 – Alıcı kredisi", 3)
    p(doc, "Türk aracı banka borçlu sıfatıyla, ihracat kredi/garanti kuruluşu destekli kredi = yurt dışından nakdi kredi.")
    p(doc, "KUKE, EKN, SACE sendikasyonda → 27. md. bağlamı.")

    h(doc, "Madde 19 – Banka aracılığı", 3)
    p(doc, "Yurt dışı nakdi döviz kredi banka aracılığıyla serbest (32 Karar 17 ile birlikte).")
    p(doc, "(7) Sözleşme + geri ödeme planı temin.")
    p(doc, "(6) Madde 14 uyum kontrolü.")
    p(doc, "(8) Geri ödeme izleme.")

    h(doc, "Madde 14 – Genel kredi kuralları", 3)
    p(doc, "Döviz geliri, 15 milyon USD limit vb. Kredi kullanan Türkiye'de yerleşik kişi için.")
    p(doc, "Eximbank banka olarak 21(a) istisnasında döviz geliri şartı aranmaz; yine 14/19 kontrol rejimi.")

    h(doc, "Madde 23 – Paranın yolu", 3)
    p(doc, "Yurt dışı kredi bedeli aracı bankaya gelmeli esastır.")
    p(doc, "Kredi kullanan haricindeki TR hesaplarına gönderilmemeli (faktoring istisnası hariç).")
    p(doc, "Bu işte: Eximbank → SCB (ajan); SCB → ihracatçı. Aracı bypass yok; 30. md. sendikasyon kullandırımı.")

    h(doc, "Madde 22 (b) – Not", 3)
    p(doc, "ECA kaynaklı kredinin yurt dışındaki ihracatçıya ödemesi: yurda getirme şartı yok.")
    p(doc, "Burada lehtar Türkiye'de; doğrudan bu bent değil; 30+23+27 rejimi.")

    h(doc, "Madde 17 – Risk Merkezi", 3)
    p(doc, "Yurt içi + yurt dışı TÜM döviz kredi kullanımları bildirilir.")
    p(doc, "Kullandıran veya aracılık eden banka bildirir → Eximbank.")
    p(doc, "Ajan bankaya transfer bildirimi kaldırmaz.")
    p(doc, "30. md. sadece İstatistik/dış finansman no muaf; Risk Merkezi muaf DEĞİL.")

    h(doc, "Madde 31 – Geri ödeme", 3)
    p(doc, "(3) Geri ödemede güncel bakiye Risk Merkezine.")
    p(doc, "(4) Kullandırma ve geri ödeme bankası farklıysa yazılı bildirim.")

    h(doc, "Madde 33 – Bildirim tarihi", 3)
    p(doc, "İhracat kredi kurumu/garanti kuruluşu kaynaklı: sözleşme imzası / gayrinakdi tesis tarihi.")
    p(doc, "Her sendikasyon otomatik değil; ECA kaynaklı işlemde.")

    h(doc, "Madde 25 – Not (bu iş için genelde değil)", 3)
    p(doc, "Türkiye Cumhuriyeti adına Bakanlık borçlu/garantör kredileri; 23. md. aranmayabilir.")
    p(doc, "Tanzanya devleti alıcı kredisi ≠ otomatik Madde 25; Hazine borçlu değilse 25 yok.")

    h(doc, "32 sayılı Karar Madde 17", 3)
    p(doc, "Nakdi döviz kredi banka aracılığıyla kullanılabilir.")

    h(doc, "2017/4 Tebliğ – İhracat taahhüdü (KRİTİK)", 3)
    p(doc, "Taahhüt = belgeli/belgesiz İHRACAT kredileri için vergi istisnası amacıyla ihracat/döviz getirme yükümlülüğü.")
    p(doc, "Alıcı kredisinde ihracatçı KREDİ KULLANMIYOR; satış bedeli alıyor.")
    p(doc, "SONUÇ: 2017/4 taahhüdü doğmaz; taahhüt kapatma uygulanmaz.")
    p(doc, "Eximbank programında gümrük beyannamesi = program ihracat teyidi; taahhüt kapatma değil.")

    # I
    h(doc, "I. RİSK MERKEZİ vs İSTATİSTİK")
    t(doc, ["", "Risk Merkezi (17.md)", "İstatistik GM (30.md)"], [
        ("Ne", "Bankalar Birliği kredi takibi", "MB dış borç istatistiği"),
        ("Sendikasyon katılım payı", "BİLDİRİM VAR", "BİLDİRİM YOK"),
        ("Kim", "Kullandıran/aracı banka (Eximbank)", "—"),
    ])

    # J
    h(doc, "J. MANTIK ZİNCİRİ (baskıda sırayla)")
    b(doc, "1. Alıcı kredisi + sendikasyon → 27. ve 30. madde.")
    b(doc, "2. Eximbank katılımcı; Standard Chartered ajan (mail metni).")
    b(doc, "3. Ödeme: katılımcı → ajan → ihracatçı TR; 30. md. sendikasyon dağıtımı.")
    b(doc, "4. 23. md.: aracı atlanmıyor.")
    b(doc, "5. Taahhüt doğmaz (2017/4 / alıcı kredisi).")
    b(doc, "6. Açık yasak yok → engel teşkil etmez.")
    b(doc, "7. Risk Merkezi (17) ve sözleşme takibi (19) uygulanır.")

    # K
    h(doc, "K. SIK SORULAR – BASKI CEVAPLARI")
    t(doc, ["Soru", "Cevap"], [
        ("Ajan Eximbank mı?", "Hayır. Ajan = Standard Chartered. Eximbank katılımcı, payı ajan'a gönderir."),
        ("Taahhüt var mı?", "2017/4 ihracat taahhüdü alıcı kredisinde doğmaz. İhracatçı kredi kullanmıyor."),
        ("Risk Merkezi?", "Evet, 17. madde. Ajan transferi muaf tutmaz."),
        ("İstatistik bildirimi?", "Katılım payı için dış borç bildirimi yok (30.md). Risk Merkezi ayrı."),
        ("SCB ne?", "Standard Chartered Bank, ajan banka."),
        ("Engel yok emin misin?", "SHG+32 Karar'da bu akışı yasaklayan açık hüküm yok. Sözleşme ile uyum dosyada teyit."),
        ("Eximbank programı?", "Program ihracat teyidi isteyebilir; bu taahhüt kapatma değil."),
        ("23. madde ihracatçıya?", "Üçüncü kişi kısıtı; burada ajan anlaşmaya göre ödüyor, katılımcı ajan'a aktarıyor."),
    ])

    # L
    h(doc, "L. KULLANDIRIM ÖNCESİ YÜKÜMLÜLÜKLER (nihai mail)")
    b(doc, "Kredi sözleşmesi + geri ödeme planı (19(7)); izleme (19(8)).")
    b(doc, "Madde 14 + 19(6) uyum kontrolü.")
    b(doc, "Risk Merkezi bildirimi (17).")
    b(doc, "Katılım payı: dış finansman no alınmaz; İstatistik dış borç bildirimi yapılmaz (30).")
    b(doc, "ECA kaynaklı işlemde bildirim tarihi (33).")
    b(doc, "Kullandırma/geri ödeme bankası farklıysa 31(4).")
    b(doc, "İhracat taahhüdü kapatma — UYGULANMAZ.")

    # M
    h(doc, "M. GÖNDERMEDEN ÖNCE TEYİT (işlem birimine 3 soru)")
    b(doc, "Anlaşmada ajan banka Standard Chartered mı?")
    b(doc, "Akış Eximbank → SCB → ihracatçı TR hesabı mı?")
    b(doc, "Eximbank anlaşmada katılımcı / aracı borçlu sıfatında mı?")

    # N
    h(doc, "N. KAYNAKLAR")
    b(doc, "Sermaye Hareketleri Genelgesi (TCMB sitesi, güncel PDF).")
    b(doc, "32 sayılı Karar Madde 17.")
    b(doc, "2017/4 sayılı Tebliğ (ihracat taahhüdü tanımı).")
    b(doc, "Gelen iç mail (işlem tanımı).")

    # O
    h(doc, "O. DOSYADA BULUNDUR")
    b(doc, "Gelen orijinal mail.")
    b(doc, "Nihai görüş maili (taahhüt takibi olmayan versiyon).")
    b(doc, "Bu çalışma kağıdı.")
    b(doc, "Kredi anlaşması özeti (borçlu, lehtar, ajan maddeleri).")
    b(doc, "Sermaye Hareketleri Genelgesi PDF.")

    # P
    h(doc, "P. ELEVATOR PITCH (30 saniye)")
    p(doc, "Eximbank sendikasyona katılıyor; kullandırımda pay Standard Chartered'a (ajan) gidiyor, ajan Yapı Merkezi'ne ödüyor. Genelge 30. madde sendikasyon dağıtımını düzenliyor. Alıcı kredisinde ihracat taahhüdü doğmaz. Mevzuatta engel yok; Risk Merkezi ve sözleşme takibi var.")

    # Q
    h(doc, "Q. NEYE GÜVENİYORUZ / NE SÖZLEŞMEYE BAĞLI")
    t(doc, ["Konu", "Durum"], [
        ("Engel yok görüşü", "SHG 30,27,19,23 + 32 Karar — güvenli"),
        ("Taahhüt yok", "2017/4 tanımı + alıcı kredisi mantığı — güvenli"),
        ("Eximbank katılımcı, SCB ajan", "Gelen mail metni — güvenli"),
        ("Madde 27 nitelendirme", "Anlaşmada Eximbank aracı borçlu yapısı — sözleşmeyle teyit"),
        ("Madde 30(1) tam uyum", "Krediyi kullanan TR yerleşik yapısı anlaşmada — teyit"),
    ])

    doc.save(OUT)
    print(f"Oluşturuldu: {OUT}")


if __name__ == "__main__":
    build()
